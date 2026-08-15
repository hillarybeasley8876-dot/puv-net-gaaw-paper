# -*- coding: utf-8 -*-
"""章节写入（模板样式 + OMML 公式版），替代 docx_emit.py。

相对旧版的四处实质改动：
  1. 公式一律走 OMML（docx_base.append_omml），不再做纯文本降级——
     用户意见第 1 条即由降级导致（`pii=₁N`、`\\mathcalP` 字符本身就错）。
  2. 段落样式改用模板自带的 Heading 1/2/3、Caption、List Paragraph，
     不再手工设字号，避免与模板不一致。
  3. 图题去重（任务 C）：正文若已有「图 X.Y …」的独立说明行，
     且该图在本章有图下 Caption，则跳过正文那一行，只保留 Caption。
  4. 表题在表上、图题在图下（模板惯例：15 处图下 vs 7 处图上）。
"""
import io
import os
import re

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Cm, Pt

from docx_base import (BODY_SIZE, CAP_INDENT, CAP_SIZE, EA_FONT,
                       EN_FONT, IMG_BASES, NOTE_SIZE,
                       ROOT, S_CAP, S_EQ, S_H1, S_H2, S_H3, TAB_SIZE,
                       add_rich, append_omml, clear_char_indent, para,
                       set_first_line_indent, set_run_font, style_names)

FIG_DECL_HDR = re.compile(r'^\|\s*(?:图编号|编号)\s*\|')
FIG_ROW = re.compile(r'^图\s*(\d+\.\d+)')
TAB_TITLE = re.compile(r'^\*\*表\s*(\d+\.\d+)\s+(.+?)\*\*\s*$')
H1 = re.compile(r'^#\s+(.+)$')
H2 = re.compile(r'^##\s+(\d+\.\d+)\s+(.+)$')
H3 = re.compile(r'^###\s+(\d+\.\d+\.\d+)\s+(.+)$')
# 正文里「图 X.Y 给出/显示/如…」的独立说明段（任务 C 去重对象之一）
FIG_TEXT_LINE = re.compile(r'^图\s*(\d+\.\d+)\s')


def resolve_img(rel):
    for b in IMG_BASES:
        p = os.path.join(ROOT, b, rel) if b else os.path.join(ROOT, rel)
        if os.path.exists(p):
            return p
    return None


def parse_table_block(lines, i):
    """按 markdown 表格切分单元格。

    注意：单元格里可能出现 LaTeX 转义竖线 `\\|`（如表 6.3 表头
    `$\\|\\Delta\\| / \\mathrm{SE}$`）。若直接 split('|')，该公式会被
    切成三段、行内公式识别失效、LaTeX 源码漏进正文。故先占位保护。
    """
    SENT = '\x00PIPE\x00'
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith('|'):
        raw = lines[i].strip().replace('\\|', SENT)
        cells = [c.strip().replace(SENT, '\\|')
                 for c in raw.strip('|').split('|')]
        if not all(re.fullmatch(r':?-{2,}:?', c) for c in cells if c):
            rows.append(cells)
        i += 1
    return rows, i


def collect_fig_numbers(lines):
    """预扫本章的图声明编号集合，供图题去重判断。"""
    nums = set()
    for ln in lines:
        m = re.match(r'^\|\s*图\s*(\d+\.\d+)\s*\|', ln.strip())
        if m:
            nums.add(m.group(1))
    return nums


def add_caption(doc, text):
    """图题/表题：模板 Caption 样式（居中，首行 21pt 为模板实测惯例）。"""
    p = para(doc, S_CAP, body=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_char_indent(p)
    set_first_line_indent(p, CAP_INDENT)
    add_rich(p, text, size=CAP_SIZE)
    return p


def add_image(doc, path, width_cm=13.5):
    p = para(doc, body=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent(p, Pt(0))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(path, width=Cm(width_cm))
    return p


def add_md_table(doc, rows):
    if not rows:
        return None
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ''
            cell = cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_first_line_indent(p, Pt(0))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_rich(p, txt, size=TAB_SIZE, base_bold=(ri == 0))
    return t


def emit_chapter(doc, path, chapter_title, stats, dedup_log=None):
    lines = io.open(path, encoding='utf-8').read().split('\n')
    fig_nums = collect_fig_numbers(lines)
    chap = os.path.basename(path)

    p = para(doc, S_H1)
    r = p.add_run(chapter_title)
    if S_H1 not in style_names(doc):
        set_run_font(r, '黑体', '黑体', 16, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats['chapter'] += 1

    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s or H1.match(s) or s.startswith('>') or s.startswith('---'):
            i += 1
            continue

        m = H2.match(s)
        if m:
            p = para(doc, S_H2)
            r = p.add_run(f'{m.group(1)} {m.group(2)}')
            if S_H2 not in style_names(doc):
                set_run_font(r, '黑体', '黑体', 15, False)
            stats['h2'] += 1
            i += 1
            continue

        m = H3.match(s)
        if m:
            p = para(doc, S_H3)
            r = p.add_run(f'{m.group(1)} {m.group(2)}')
            if S_H3 not in style_names(doc):
                set_run_font(r, '黑体', '黑体', 14, False)
            stats['h3'] += 1
            i += 1
            continue

        # 独立公式块 -> OMML（居中，模板公式样式）
        if s == '$$':
            body, tag = [], None
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                t2 = lines[i].strip()
                mt = re.search(r'\\tag\{([\d.]+)\}', t2)
                if mt:
                    tag = mt.group(1)
                    t2 = re.sub(r'\\tag\{[\d.]+\}', '', t2).strip()
                if t2:
                    body.append(t2)
                i += 1
            i += 1
            latex = ' '.join(body)
            sty = S_EQ if S_EQ in style_names(doc) else None
            p = para(doc, sty, body=False)
            pf = p.paragraph_format
            # 模板 MTDisplayEquation 样式带 199.5pt 首行缩进 + JUSTIFY，
            # 直接套用会把公式推到版心右侧再折行（实测 p088：公式首行
            # 从 x0=386pt 起笔、总宽仅 119pt 却折成两行）。故显式清零缩进。
            set_first_line_indent(p, Pt(0))
            pf.left_indent = Pt(0)
            pf.right_indent = Pt(0)
            clear_char_indent(p)   # 关键：清 firstLineChars=1900
            if tag:
                # 「公式居中 + 编号右对齐」在 Word 中的标准做法：
                # 段落左对齐，前置一个居中制表位放公式，末尾一个右对齐
                # 制表位放编号。段落若设 CENTER，右对齐制表位会失效。
                # 版心宽 = 210.1mm - 2×31.7mm = 146.7mm ≈ 415.9pt。
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.tab_stops.add_tab_stop(Pt(207.9),
                                          WD_TAB_ALIGNMENT.CENTER)
                pf.tab_stops.add_tab_stop(Pt(415.9), WD_TAB_ALIGNMENT.RIGHT)
                p.add_run('\t')
                append_omml(p, latex, display=False)
                r = p.add_run('\t（' + tag + '）')
                set_run_font(r, EA_FONT, EN_FONT, BODY_SIZE)
            else:
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                append_omml(p, latex, display=False)
            stats['eq'] += 1
            continue

        # 图声明表 -> 插图 + 图下 Caption
        if FIG_DECL_HDR.match(s):
            rows, ni = parse_table_block(lines, i)
            fig_rows = [r0 for r0 in rows[1:]
                        if r0 and FIG_ROW.match(r0[0].strip())]
            if not fig_rows:
                add_md_table(doc, rows)
                stats['tab_bare'] += 1
                i = ni
                continue
            for row in fig_rows:
                no = re.search(r'(\d+\.\d+)', row[0]).group(1)
                title = row[1] if len(row) > 1 else ''
                png = None
                for c in row:
                    mm = re.search(r'`([^`]+\.png)`', c)
                    if mm:
                        png = mm.group(1)
                        break
                fp = resolve_img(png) if png else None
                if fp:
                    add_image(doc, fp)
                    add_caption(doc, f'图{no}  {title}')
                    stats['fig'] += 1
                else:
                    add_caption(doc, f'图{no}  {title}')
                    stats['fig_missing'] += 1
            i = ni
            continue

        # 表题在表上
        m = TAB_TITLE.match(s)
        if m:
            add_caption(doc, f'表{m.group(1)}  {m.group(2)}')
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].lstrip().startswith('|'):
                rows, nj = parse_table_block(lines, j)
                add_md_table(doc, rows)
                stats['tab'] += 1
                i = nj
            else:
                i += 1
            continue

        if s.startswith('|'):
            rows, ni = parse_table_block(lines, i)
            add_md_table(doc, rows)
            stats['tab_bare'] += 1
            i = ni
            continue

        if s.startswith('```'):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1
            for b in buf:
                p = para(doc, body=False)
                set_first_line_indent(p, Pt(0))
                p.paragraph_format.line_spacing = Pt(13)
                set_run_font(p.add_run(b), '宋体', 'Consolas', 9)
            stats['code'] += 1
            continue

        mlist = re.match(r'^(\d+\.|[-*])\s+(.+)$', s)
        if mlist:
            p = para(doc)
            set_first_line_indent(p, Pt(0))
            p.paragraph_format.left_indent = Pt(21)
            prefix = mlist.group(1)
            add_rich(p, f'{"·" if prefix in ("-", "*") else prefix} '
                        f'{mlist.group(2)}', size=BODY_SIZE)
            stats['list'] += 1
            i += 1
            continue

        if s.startswith(('表注：', '注：')):
            p = para(doc, body=False)
            set_first_line_indent(p, Pt(0))
            add_rich(p, s, size=NOTE_SIZE)
            stats['note'] += 1
            i += 1
            continue

        # 任务 C：正文重复图题去重
        mf = FIG_TEXT_LINE.match(s)
        if mf and mf.group(1) in fig_nums:
            # 该行以「图 X.Y」开头且本章确有该图的 Caption。
            # 若整行只是命名（短、无动词性说明），删除；
            # 若是分析性说明（长），保留但去掉开头的重复编号。
            no = mf.group(1)
            if len(s) <= 30:
                if dedup_log is not None:
                    dedup_log.append((chap, no, 'removed', s[:40]))
                stats['fig_dedup'] += 1
                i += 1
                continue

        p = para(doc)
        add_rich(p, s, size=BODY_SIZE)
        stats['body'] += 1
        i += 1
