# -*- coding: utf-8 -*-
"""从正稿 Markdown 生成同济格式 Word 文稿（python-docx）。

为什么不复用 scripts/build_thesis_docx.js：
  实测该链读的是 refs/codex_n11_deprecated/ 下已作废的 n=11 稿文件名
  （ch3_analysis_framework / ch4_research_design / ch5_gaaw_method / ch6_results），
  且其资产提取规则要求「markdown ![]() 图片 + 行首图题」「章内编号从 1 连续」，
  与正稿的「表格行声明图 + 加粗表题」写法完全不匹配（实测提取到图 0 表 0）。
  改造成本高于重写，故另写本脚本，直接对齐 FORMAT_TONGJI.md 的字体字号表。

格式依据（docs/FORMAT_TONGJI.md，每条可回溯官方模板 dump 段落号）：
  章标题「第N章 XXX」   黑体 16pt 粗 居中                  [157]
  一级节「N.M XXX」     黑体 15pt 非粗 左                   [158]
  二级节「N.M.K XXX」   黑体 14pt 非粗 左                   [167]
  正文                  宋体 12pt 两端 行距固定20pt 首行缩进24pt  [159]
  段内小标题「1. XXX」  宋体 12pt 左 无首行缩进               [183]
  图题「图N.M XXX」     宋体 10.5pt 居中 单倍（图下方）        [188]
  表题「表N.M XXX」     宋体 10.5pt 居中（表上方）             [192]
  公式编号「（N.M）」   宋体 12pt 右对齐                      [185]
  参考文献条目          宋体 10.5pt 行距固定16pt 悬挂21pt      [203]
  一级标题（摘要等）    黑体 16pt 粗 居中                     [115]
  摘要正文              宋体 12pt 行距20pt 首行24pt           [116]
  关键词行              宋体 12pt 粗                          [123]
  致谢正文              仿宋 12pt 行距20pt 首行24pt           [215]

注意：模板实测图题为「图3.2 标题」——编号与文字间无空格；
章标题为「第1章 引言」，「章」与标题间一个空格。以模板为准。
"""
import io
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# stdout 不重包装：与 PowerShell 的 `>` 重定向叠加会导致
# "I/O operation on closed file"。由外部 PYTHONUTF8=1 保证中文输出。
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CH = os.path.join(ROOT, 'docs', 'chapters')
OUT_DIR = os.path.join(ROOT, 'outputs', 'thesis')

TITLE_ZH = '基于梯度自适应对抗约束的点云上采样方法研究'
TITLE_EN = ('Research on Point Cloud Upsampling with '
            'Gradient-Adaptive Adversarial Constraints')

CHAPTERS = [
    ('ch1_introduction.md', '第1章 绪论'),
    ('ch2_related_work.md', '第2章 文献综述'),
    ('ch3_baseline.md', '第3章 基线模型复现与瓶颈诊断'),
    ('ch4_design.md', '第4章 模型假设和研究设计'),
    ('ch5_mechanism.md', '第5章 梯度自适应对抗权重机制'),
    ('ch6_experiments.md', '第6章 实验结果与边界分析'),
    ('ch7_conclusion.md', '第7章 结论与展望'),
]

# 图路径可能相对仓库根或 docs/
IMG_BASES = ['', 'docs']


# ----------------------------------------------------------------- 底层工具
def set_run_font(run, ea, en, size, bold=False):
    """设置中西文字体。中文必须走 w:eastAsia，否则 Word 回退到默认宋体。"""
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), ea)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)


def set_para(p, align=None, line_pt=None, first_indent_pt=None,
             hanging_pt=None, space_before=0, space_after=0):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if line_pt is not None:
        # 固定行距：exact
        pf.line_spacing = Pt(line_pt)
    if first_indent_pt is not None:
        pf.first_line_indent = Pt(first_indent_pt)
    if hanging_pt is not None:
        pf.first_line_indent = Pt(-hanging_pt)
        pf.left_indent = Pt(hanging_pt)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p


def add_para(doc, text, ea='宋体', en='宋体', size=12, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt=20,
             first_indent_pt=24, hanging_pt=None):
    p = doc.add_paragraph()
    set_para(p, align=align, line_pt=line_pt,
             first_indent_pt=first_indent_pt, hanging_pt=hanging_pt)
    if text:
        r = p.add_run(text)
        set_run_font(r, ea, en, size, bold)
    return p


# 行内标记：**粗体**、`代码`、$公式$（公式按纯文本降级，保留可读性）
INLINE = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\$[^$]+\$)')


def unescape_text(s):
    """剥纯文本里的 LaTeX 转义。

    正稿中 `6.43\\%`、`Q_4/Q_1` 这类写法出现在 $...$ **之外**，
    latex_to_text() 不经过它们（实测 p005/p008 出现 `91.3\\%` 残留）。
    这里只处理确定安全的转义，不做公式解析。
    """
    for a, b in (('\\%', '%'), ('\\&', '&'), ('\\#', '#'),
                 ('\\_', '_'), ('\\$', '$'), ('\\{', '{'), ('\\}', '}')):
        s = s.replace(a, b)
    return s


CITE_RE = re.compile(r'(\[\d+(?:,\d+)*\])')


def add_rich(p, text, ea='宋体', en='宋体', size=12, base_bold=False):
    """按行内标记分段写 run，保留粗体/等宽/公式文本。

    引用编号 [N] 按规范 §2.3 排为**上标**（模板 [165] 实证）。
    纯文本段还要剥 `\\%` 这类 LaTeX 转义（它们在 $...$ 之外）。
    """
    def emit_plain(txt, bold):
        """纯文本再按引用编号切分，[N] 走上标。"""
        for piece in CITE_RE.split(txt):
            if not piece:
                continue
            r = p.add_run(unescape_text(piece))
            set_run_font(r, ea, en, size, bold)
            if CITE_RE.fullmatch(piece):
                r.font.superscript = True

    for seg in INLINE.split(text):
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            emit_plain(seg[2:-2], True)
        elif seg.startswith('`') and seg.endswith('`'):
            r = p.add_run(seg[1:-1])
            set_run_font(r, ea, 'Consolas', size - 0.5, base_bold)
        elif seg.startswith('$') and seg.endswith('$'):
            r = p.add_run(latex_to_text(seg[1:-1]))
            set_run_font(r, ea, 'Cambria Math', size, base_bold)
        else:
            emit_plain(seg, base_bold)


GREEK = {
    'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'epsilon': 'ε',
    'theta': 'θ', 'lambda': 'λ', 'mu': 'μ', 'rho': 'ρ', 'sigma': 'σ',
    'tau': 'τ', 'phi': 'φ', 'psi': 'ψ', 'omega': 'ω', 'Delta': 'Δ',
    'Sigma': 'Σ', 'Omega': 'Ω', 'nabla': '∇', 'eta': 'η', 'zeta': 'ζ',
    'chi': 'χ', 'kappa': 'κ', 'xi': 'ξ',
}
SUB = str.maketrans('0123456789+-=()aeoxhklmnpst',
                    '₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₒₓₕₖₗₘₙₚₛₜ')
SUP = str.maketrans('0123456789+-=()n', '⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ')


def latex_to_text(s):
    """把行内 LaTeX 降级为可读 Unicode 文本。

    这不是完整渲染器：目的是让 Word 稿在没有公式编辑器的情况下仍可读，
    独立公式另走 add_equation()。已知取舍：分式写作 a/b。

    覆盖范围按全文实测频次确定（scripts/_tmp/audit_latex_cmds.py 枚举出
    39 种命令）。首版只处理 mathrm/mathbf/text，漏掉 mathcal/mathbb/dots/
    lVert 等，导致 p026 出现大面积 `\\mathcalP`、`\\mathbbR` 残留 ——
    这类残留必须靠视觉验收发现，代码层的「跑通」不代表渲染正确。
    """
    # 1) 字体/样式命令：剥命令名，保留花括号内容（可嵌套两层）
    for _ in range(3):
        s = re.sub(r'\\(?:mathrm|mathbf|mathcal|mathbb|mathit|mathsf|text|'
                   r'operatorname|boldsymbol)\s*\{([^{}]*)\}', r'\1', s)
    # 2) 结构命令
    s = re.sub(r'\\frac\s*\{([^{}]*)\}\s*\{([^{}]*)\}', r'(\1)/(\2)', s)
    s = re.sub(r'\\sqrt\s*\{([^{}]*)\}', r'√(\1)', s)
    s = re.sub(r'\\(?:hat|bar|tilde|vec|dot)\s*\{([^{}]*)\}', r'\1', s)
    s = re.sub(r'\\(?:hat|bar|tilde|vec|dot)\s+', '', s)
    # 3) 希腊字母与运算符
    for k, v in GREEK.items():
        s = s.replace('\\' + k, v)
    for a, b in (
            ('\\times', '×'), ('\\cdot', '·'), ('\\leq', '≤'),
            ('\\geq', '≥'), ('\\approx', '≈'), ('\\neq', '≠'),
            ('\\to', '→'), ('\\triangleq', '≜'), ('\\in', '∈'),
            ('\\subset', '⊂'), ('\\subseteq', '⊆'), ('\\sum', '∑'),
            ('\\prod', '∏'), ('\\infty', '∞'), ('\\pm', '±'),
            ('\\nabla', '∇'), ('\\partial', '∂'), ('\\ll', '≪'),
            ('\\gg', '≫'), ('\\sim', '~'), ('\\propto', '∝'),
            ('\\forall', '∀'), ('\\exists', '∃'), ('\\dots', '…'),
            ('\\ldots', '…'), ('\\cdots', '⋯'),
            ('\\lVert', '‖'), ('\\rVert', '‖'),
            ('\\lvert', '|'), ('\\rvert', '|'),
            ('\\langle', '⟨'), ('\\rangle', '⟩'),
            ('\\min', 'min'), ('\\max', 'max'),
            ('\\inf', 'inf'), ('\\sup', 'sup'),
            ('\\log', 'log'), ('\\exp', 'exp'),
            ('\\bigl', ''), ('\\bigr', ''), ('\\Bigl', ''),
            ('\\Bigr', ''), ('\\left', ''), ('\\right', ''),
            ('\\,', ' '), ('\\;', ' '), ('\\:', ' '),
            ('\\quad', '  '), ('\\qquad', '    '),
            ('\\%', '%'), ('\\&', '&'), ('\\#', '#'),
            ('\\{', '{'), ('\\}', '}'), ('\\|', '‖'),
    ):
        s = s.replace(a, b)
    # 4) 上下标（花括号形式优先，再处理单字符形式）
    s = re.sub(r'_\{([^{}]*)\}', lambda m: m.group(1).translate(SUB), s)
    s = re.sub(r'\^\{([^{}]*)\}', lambda m: m.group(1).translate(SUP), s)
    s = re.sub(r'_(\w)', lambda m: m.group(1).translate(SUB), s)
    s = re.sub(r'\^(\w)', lambda m: m.group(1).translate(SUP), s)
    # 5) 收尾：剥残留花括号与换行符
    s = s.replace('\\\\', ' ').replace('{', '').replace('}', '')
    # 未识别的命令：去掉反斜杠保留词干，避免出现 \foo 这类原样残留
    s = re.sub(r'\\([a-zA-Z]+)', r'\1', s)
    return re.sub(r'\s{3,}', '  ', s).strip()


def add_equation(doc, body, tag):
    """公式独占段 + 右对齐全角编号（规范禁止半角括号）。"""
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=None,
             first_indent_pt=0, space_before=6, space_after=6)
    r = p.add_run(latex_to_text(body))
    set_run_font(r, '宋体', 'Cambria Math', 12)
    if tag:
        r2 = p.add_run('\t（' + tag + '）')
        set_run_font(r2, '宋体', '宋体', 12)
    return p


def add_image(doc, path, width_cm=14.0):
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
             space_before=6, space_after=3)
    p.add_run().add_picture(path, width=Cm(width_cm))
    return p


def add_caption(doc, text, size=10.5):
    """图题/表题：宋体 10.5pt 居中，无首行缩进。"""
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=None,
             first_indent_pt=0, space_before=3, space_after=6)
    add_rich(p, text, size=size)
    return p


def add_md_table(doc, rows):
    """写 markdown 表格。首行表头加粗、居中；全表 10.5pt。"""
    if not rows:
        return None
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ''
            cell = cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=None,
                     first_indent_pt=0)
            add_rich(p, txt, size=10.5, base_bold=(ri == 0))
    return t
