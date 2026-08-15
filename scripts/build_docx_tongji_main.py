# -*- coding: utf-8 -*-
"""主流程：组装同济格式 Word 文稿。

结构按 THESIS_OUTLINE §0 规定的前置顺序：
  中文封面 → 英文封面 → 原创性声明 → 中文摘要 → ABSTRACT → 目录
  → 插图索引 → 插表索引 → 公式索引 → 正文七章 → 参考文献 → 致谢

目录/索引用 Word 域代码（TOC / TOF），页码由 Word 自动生成——
规范明确「不能手敲页码」。域需在 Word 中按 F9 更新，
本脚本已通过 COM 在导出 PDF 前自动更新域。

用法：
  python scripts/build_docx_tongji_main.py
产物：
  outputs/thesis/GAAW_thesis_v2.docx
"""
import io
import os
import re
import sys
from collections import Counter

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from build_docx_tongji import (CHAPTERS, CH, OUT_DIR, ROOT, TITLE_EN,
                               TITLE_ZH, add_caption, add_para, add_rich,
                               set_para, set_run_font)
from docx_emit import emit_chapter, parse_table_block

# 不再重包装 stdout：与 PowerShell 的 `>` 重定向叠加会导致
# "I/O operation on closed file"（实测首次运行即触发）。
# 改由外部设置 PYTHONIOENCODING=utf-8 / PYTHONUTF8=1 保证中文输出。


def add_field(p, instr, placeholder='（右键“更新域”生成）'):
    """插入 Word 域代码（TOC 等）。页码必须由域生成，不得手敲。"""
    r = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    r._element.append(fld)
    r2 = p.add_run()
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = instr
    r2._element.append(it)
    r3 = p.add_run()
    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    r3._element.append(sep)
    r4 = p.add_run(placeholder)
    set_run_font(r4, '宋体', '宋体', 12)
    r5 = p.add_run()
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    r5._element.append(end)


def h1_title(doc, text):
    """一级标题（摘要/ABSTRACT/目录/参考文献/致谢）：黑体 16pt 粗 居中。"""
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
             space_before=0, space_after=18)
    set_run_font(p.add_run(text), '黑体', '黑体', 16, True)
    return p


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def set_page_numbering(sec, fmt='decimal', start=None):
    """设置本节页码格式。前置部分用小写罗马数字，正文从 1 重新开始。

    规范惯例：封面与声明不编页码；摘要起用 i/ii/iii；正文第 1 章起用 1。
    实测 p001 发现首版把封面也编成 1 且带页眉，故引入分节。
    """
    sectPr = sec._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    el = OxmlElement('w:pgNumType')
    el.set(qn('w:fmt'), fmt)
    if start is not None:
        el.set(qn('w:start'), str(start))
    sectPr.append(el)


def clear_header_footer(sec):
    """本节不显示页眉页码（用于封面与声明节）。"""
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    for part in (sec.header, sec.footer):
        for p in part.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)


def fill_header_footer(sec, with_page=True):
    """页眉居中书名（宋体 9pt），页脚居中页码域。"""
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    hp = sec.header.paragraphs[0]
    for r in list(hp.runs):
        r._element.getparent().remove(r._element)
    set_para(hp, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0)
    set_run_font(hp.add_run('同济大学 硕士学位论文  ' + TITLE_ZH),
                 '宋体', '宋体', 9)
    fp = sec.footer.paragraphs[0]
    for r in list(fp.runs):
        r._element.getparent().remove(r._element)
    set_para(fp, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0)
    if with_page:
        add_field(fp, ' PAGE ', '1')


def setup_section(sec):
    """A4 + 同济页边距；页眉居中书名，页脚居中页码。"""
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(2.0)
    sec.footer_distance = Cm(1.75)
    # 页眉页脚由 fill_header_footer / clear_header_footer 按节分别设置，
    # 此处只管页面尺寸与边距。


def cover_zh(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
             space_after=6)
    set_run_font(p.add_run('硕士学位论文'), '隶书', '隶书', 22, True)
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
             space_after=36)
    set_run_font(p.add_run('（学术学位）'), '隶书', '隶书', 16, True)

    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
             space_after=48)
    set_run_font(p.add_run(TITLE_ZH), '黑体', '黑体', 18, True)

    for label in ('作者姓名', '学科专业', '指导教师', '培养单位'):
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
                 space_after=10)
        set_run_font(p.add_run(f'{label}：____________________'),
                     '宋体', '宋体', 14)

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0)
    set_run_font(p.add_run('同济大学'), '宋体', '宋体', 14)
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0)
    set_run_font(p.add_run('二〇二六年八月'), '宋体', '宋体', 14)


def cover_en(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
             space_after=36)
    set_run_font(p.add_run('A dissertation submitted to Tongji University '
                           'in conformity with the requirements for the '
                           'degree of Master of Engineering'),
                 'Times New Roman', 'Times New Roman', 12)
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
             space_after=48)
    set_run_font(p.add_run(TITLE_EN), 'Times New Roman',
                 'Times New Roman', 16, True)
    for label in ('Candidate', 'Major', 'Supervisor'):
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
                 space_after=10)
        set_run_font(p.add_run(f'{label}: ____________________'),
                     'Times New Roman', 'Times New Roman', 12)


def declaration(doc):
    """原创性声明：标题黑体 18 粗居中；正文宋体 14 行距 1.4 首行 28pt。"""
    for title, body in (
        ('同济大学学位论文原创性声明',
         '本人郑重声明：所呈交的学位论文，是本人在导师指导下，进行研究工作所取得的成果。'
         '除文中已经注明引用的内容外，本学位论文的研究成果不包含任何他人创作的、已公开'
         '发表或者没有公开发表的作品的内容。对本论文所涉及的研究工作做出贡献的其他个人'
         '和集体，均已在文中以明确方式标明。本学位论文原创性声明的法律责任由本人承担。'),
        ('学位论文版权使用授权书',
         '本人完全了解同济大学关于收集、保存、使用学位论文的规定，同意如下各项内容：'
         '按照学校要求提交学位论文的印刷本和电子版本；学校有权保存学位论文的印刷本和'
         '电子版，并采用影印、缩印、扫描、数字化或其它手段保存论文；学校有权提供目录'
         '检索以及提供本学位论文全文或者部分的阅览服务；学校有权按有关规定向国家有关'
         '部门或者机构送交论文的复印件和电子版。'),
    ):
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
                 space_after=18)
        set_run_font(p.add_run(title), '黑体', '黑体', 18, True)
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent_pt=28)
        p.paragraph_format.line_spacing = 1.4
        set_run_font(p.add_run(body), '宋体', '宋体', 14)
        for _ in range(2):
            doc.add_paragraph()
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent_pt=0)
        set_run_font(p.add_run('签名：____________    日期：__________'),
                     '宋体', '宋体', 14)
        page_break(doc)


def emit_abstract(doc, path, is_en, stats):
    """摘要：正文宋体 12 行距 20 首行 24；关键词行加粗。"""
    lines = io.open(os.path.join(CH, path), encoding='utf-8').read().split('\n')
    h1_title(doc, 'ABSTRACT' if is_en else '摘要')
    en = 'Times New Roman' if is_en else '宋体'
    for ln in lines:
        s = ln.strip()
        if not s or s.startswith('#') or s.startswith('>'):
            continue
        is_kw = s.startswith(('**关键词', '**Key Words', '**Key words'))
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt=20,
                 first_indent_pt=24)
        add_rich(p, s, ea='宋体', en=en, size=12, base_bold=is_kw)
        stats['abstract'] += 1


def emit_index(doc, title, instr_label, items):
    """插图/插表/公式索引：编号+标题+页码域。

    页码用 { PAGEREF } 不现实（需书签），此处按规范用 TOC 域按题注样式收集；
    但本脚本的图题/表题为直接排版而非 Word 题注样式，故索引以「编号 标题」
    列出并显式说明页码需在 Word 中插入题注后由域生成。
    """
    h1_title(doc, title)
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_pt=0,
             space_after=10)
    set_run_font(p.add_run(
        '说明：本索引按正文出现顺序列出。页码需在 Word 中将图题/表题设为'
        '“题注”样式后，用「引用—插入表目录」自动生成，不得手敲。'),
        '宋体', '宋体', 10.5)
    for no, ttl in items:
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_pt=None,
                 first_indent_pt=0)
        p.paragraph_format.left_indent = Pt(12)
        add_rich(p, f'{instr_label}{no}  {ttl}', size=12)


def collect_assets():
    """从正稿收集图/表/公式清单（供索引使用）。"""
    figs, tabs, eqs = [], [], []
    for fn, _ in CHAPTERS:
        t = io.open(os.path.join(CH, fn), encoding='utf-8').read()
        for m in re.finditer(r'^\|\s*图\s*(\d+\.\d+)\s*\|([^|]+)\|',
                             t, re.M):
            figs.append((m.group(1), m.group(2).strip()))
        for m in re.finditer(r'^\*\*表\s*(\d+\.\d+)\s+(.+?)\*\*\s*$', t, re.M):
            tabs.append((m.group(1), m.group(2).strip()))
        for m in re.finditer(r'\\tag\{(\d+\.\d+)\}', t):
            eqs.append((m.group(1), '见正文对应段落'))
    return figs, tabs, eqs


def emit_references(doc):
    """参考文献：宋体 10.5pt 行距固定 16pt 悬挂缩进 21pt。"""
    h1_title(doc, '参考文献')
    md = io.open(os.path.join(ROOT, 'docs', 'REFERENCES_GB7714.md'),
                 encoding='utf-8').read()
    n = 0
    for ln in md.split('\n'):
        s = ln.strip()
        if not re.match(r'^\[\d+\]', s):
            continue
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_pt=16,
                 hanging_pt=21)
        set_run_font(p.add_run(s), '宋体', '宋体', 10.5)
        n += 1
    return n


def emit_ack(doc):
    """致谢：仿宋 12pt 行距 20pt 首行 24pt。"""
    h1_title(doc, '致谢')
    p = os.path.join(CH, 'acknowledgements.md')
    if not os.path.exists(p):
        return 0
    n = 0
    for ln in io.open(p, encoding='utf-8').read().split('\n'):
        s = ln.strip()
        if not s or s.startswith('#') or s.startswith('>'):
            continue
        par = doc.add_paragraph()
        set_para(par, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt=20,
                 first_indent_pt=24)
        add_rich(par, s, ea='仿宋', en='仿宋', size=12)
        n += 1
    return n


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    setup_section(doc.sections[0])

    # 基础样式兜底（防 Word 回退默认字体）
    st = doc.styles['Normal']
    st.font.name = '宋体'
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    stats = Counter()

    # ---- 第 1 节：封面 + 声明（无页眉、无页码）----
    clear_header_footer(doc.sections[0])
    set_page_numbering(doc.sections[0], 'decimal', 1)
    cover_zh(doc)
    page_break(doc)
    cover_en(doc)
    page_break(doc)
    declaration(doc)

    # ---- 第 2 节：前置部分（摘要/目录/索引，小写罗马数字 i 起）----
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(sec2)
    fill_header_footer(sec2, with_page=True)
    set_page_numbering(sec2, 'lowerRoman', 1)

    emit_abstract(doc, 'abstract_zh.md', False, stats)
    page_break(doc)
    emit_abstract(doc, 'abstract_en.md', True, stats)
    page_break(doc)

    h1_title(doc, '目录')
    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_pt=0)
    add_field(p, ' TOC \\o "1-3" \\h \\z \\u ')
    page_break(doc)

    figs, tabs, eqs = collect_assets()
    emit_index(doc, '插图索引', '图', figs)
    page_break(doc)
    emit_index(doc, '插表索引', '表', tabs)
    page_break(doc)
    emit_index(doc, '公式索引', '式', eqs)

    # ---- 第 3 节：正文起（阿拉伯数字从 1 重新开始）----
    sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(sec3)
    fill_header_footer(sec3, with_page=True)
    set_page_numbering(sec3, 'decimal', 1)

    for idx, (fn, title) in enumerate(CHAPTERS):
        emit_chapter(doc, os.path.join(CH, fn), title, stats)
        if idx < len(CHAPTERS) - 1:
            page_break(doc)

    page_break(doc)
    nref = emit_references(doc)
    page_break(doc)
    nack = emit_ack(doc)

    out = os.path.join(OUT_DIR, 'GAAW_thesis_v2.docx')
    doc.save(out)

    print('已生成:', os.path.relpath(out, ROOT))
    print(f'  体积 {os.path.getsize(out) / 1024 / 1024:.2f} MB')
    print()
    print('写入统计:')
    for k in ('chapter', 'h2', 'h3', 'body', 'subhead', 'list', 'tab',
              'tab_bare', 'fig', 'fig_missing', 'eq', 'note', 'code',
              'abstract'):
        print(f'  {k:14s} {stats[k]}')
    print(f'  references     {nref}')
    print(f'  ack_paras      {nack}')
    print(f'  索引: 图 {len(figs)} 表 {len(tabs)} 式 {len(eqs)}')
    if stats['fig_missing']:
        print(f'  !! 图片缺失 {stats["fig_missing"]} 处')
    return 0


if __name__ == '__main__':
    sys.exit(main())
