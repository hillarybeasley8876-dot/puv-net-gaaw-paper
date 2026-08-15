# -*- coding: utf-8 -*-
"""主流程（模板基底 + OMML 公式版）。

产物：outputs/thesis/GAAW_thesis_v3.docx

相对 v2 的改动见 docx_base.py 与 docx_emit2.py 的模块说明。
本文件负责：前置部分、分节页码、目录/索引域、参考文献、致谢、导出自检。
"""
import io
import os
import re
import sys
import zipfile
from collections import Counter

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx_base import (BODY_SIZE, CH, CHAPTERS, EA_FONT, EN_FONT,
                       OUT_DIR, REF_HANG,
                       REF_LINE, REF_SIZE, ROOT, S_CAP, S_H1, S_REF,
                       TITLE_EN, TITLE_ZH, add_rich, append_omml, clear_char_indent,
                       new_document, para, set_first_line_indent,
                       set_run_font, style_names)
from docx_emit2 import emit_chapter


def add_field(p, instr, placeholder='（更新域后生成）'):
    r = p.add_run()
    f1 = OxmlElement('w:fldChar')
    f1.set(qn('w:fldCharType'), 'begin')
    r._element.append(f1)
    r2 = p.add_run()
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = instr
    r2._element.append(it)
    r3 = p.add_run()
    f2 = OxmlElement('w:fldChar')
    f2.set(qn('w:fldCharType'), 'separate')
    r3._element.append(f2)
    p.add_run(placeholder)
    r5 = p.add_run()
    f3 = OxmlElement('w:fldChar')
    f3.set(qn('w:fldCharType'), 'end')
    r5._element.append(f3)


def set_page_numbering(sec, fmt='decimal', start=None):
    sectPr = sec._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    el = OxmlElement('w:pgNumType')
    el.set(qn('w:fmt'), fmt)
    if start is not None:
        el.set(qn('w:start'), str(start))
    sectPr.append(el)


def clear_hf(sec):
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    for part in (sec.header, sec.footer):
        for p in part.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)


def fill_hf(sec, doc):
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    hp = sec.header.paragraphs[0]
    for r in list(hp.runs):
        r._element.getparent().remove(r._element)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(hp.add_run('同济大学硕士学位论文'), '宋体', '宋体', 9)
    fp = sec.footer.paragraphs[0]
    for r in list(fp.runs):
        r._element.getparent().remove(r._element)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(fp, ' PAGE ', '1')


def h1(doc, text):
    p = para(doc, S_H1)
    r = p.add_run(text)
    if S_H1 not in style_names(doc):
        set_run_font(r, '黑体', '黑体', 16, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def pbreak(doc):
    para(doc).add_run().add_break(WD_BREAK.PAGE)


def cover_zh(doc):
    for _ in range(3):
        para(doc)
    p = para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent(p, Pt(0))
    set_run_font(p.add_run('硕士学位论文'), '隶书', '隶书', 22, True)
    p = para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent(p, Pt(0))
    p.paragraph_format.space_after = Pt(36)
    set_run_font(p.add_run('（学术学位）'), '隶书', '隶书', 16, True)
    p = para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent(p, Pt(0))
    p.paragraph_format.space_after = Pt(48)
    set_run_font(p.add_run(TITLE_ZH), '黑体', '黑体', 18, True)
    for label in ('作者姓名', '学科专业', '指导教师', '培养单位'):
        p = para(doc)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_first_line_indent(p, Pt(0))
        set_run_font(p.add_run(f'{label}：____________________'),
                     '宋体', '宋体', 14)
    for _ in range(2):
        para(doc)
    for s in ('同济大学', '二〇二六年八月'):
        p = para(doc)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_first_line_indent(p, Pt(0))
        set_run_font(p.add_run(s), '宋体', '宋体', 14)


def cover_en(doc):
    for _ in range(4):
        para(doc)
    p = para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent(p, Pt(0))
    p.paragraph_format.space_after = Pt(36)
    set_run_font(p.add_run(
        'A dissertation submitted to Tongji University in conformity '
        'with the requirements for the degree of Master of Engineering'),
        'Times New Roman', 'Times New Roman', 12)
    p = para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_first_line_indent(p, Pt(0))
    p.paragraph_format.space_after = Pt(48)
    set_run_font(p.add_run(TITLE_EN), 'Times New Roman',
                 'Times New Roman', 16, True)
    for label in ('Candidate', 'Major', 'Supervisor'):
        p = para(doc)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_first_line_indent(p, Pt(0))
        set_run_font(p.add_run(f'{label}: ____________________'),
                     'Times New Roman', 'Times New Roman', 12)


def declaration(doc):
    for title, body in (
        ('同济大学学位论文原创性声明',
         '本人郑重声明：所呈交的学位论文，是本人在导师指导下，进行研究工作所取得的成果。'
         '除文中已经注明引用的内容外，本学位论文的研究成果不包含任何他人创作的、已公开'
         '发表或者没有公开发表的作品的内容。对本论文所涉及的研究工作做出贡献的其他个人'
         '和集体，均已在文中以明确方式标明。本学位论文原创性声明的法律责任由本人承担。'),
        ('学位论文版权使用授权书',
         '本人完全了解同济大学关于收集、保存、使用学位论文的规定，同意如下各项内容：'
         '按照学校要求提交学位论文的印刷本和电子版本；学校有权保存学位论文的印刷本和'
         '电子版，并采用影印、缩印、扫描、数字化或其它手段保存论文；学校有权提供目录'
         '检索以及提供本学位论文全文或者部分的阅览服务；学校有权按有关规定向国家有关'
         '部门或者机构送交论文的复印件和电子版。'),
    ):
        p = para(doc)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_first_line_indent(p, Pt(0))
        p.paragraph_format.space_after = Pt(18)
        set_run_font(p.add_run(title), '黑体', '黑体', 18, True)
        p = para(doc)
        set_first_line_indent(p, Pt(28))
        p.paragraph_format.line_spacing = 1.4
        set_run_font(p.add_run(body), '宋体', '宋体', 14)
        for _ in range(2):
            para(doc)
        p = para(doc)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_first_line_indent(p, Pt(0))
        set_run_font(p.add_run('签名：____________    日期：__________'),
                     '宋体', '宋体', 14)
        pbreak(doc)


def emit_abstract(doc, fn, is_en, stats):
    lines = io.open(os.path.join(CH, fn), encoding='utf-8').read().split('\n')
    h1(doc, 'ABSTRACT' if is_en else '摘要')
    en = 'Times New Roman' if is_en else '宋体'
    for ln in lines:
        s = ln.strip()
        if not s or s.startswith('#') or s.startswith('>'):
            continue
        kw = s.startswith(('**关键词', '**Key Words', '**Key words'))
        p = para(doc)
        add_rich(p, s, ea=EA_FONT, en=en, size=BODY_SIZE, base_bold=kw)
        stats['abstract'] += 1


def collect_assets():
    figs, tabs, eqs = [], [], []
    for fn, _ in CHAPTERS:
        t = io.open(os.path.join(CH, fn), encoding='utf-8').read()
        for m in re.finditer(r'^\|\s*图\s*(\d+\.\d+)\s*\|([^|]+)\|', t, re.M):
            figs.append((m.group(1), m.group(2).strip()))
        for m in re.finditer(r'^\*\*表\s*(\d+\.\d+)\s+(.+?)\*\*\s*$', t, re.M):
            tabs.append((m.group(1), m.group(2).strip()))
        for m in re.finditer(r'\\tag\{(\d+\.\d+)\}', t):
            eqs.append((m.group(1), ''))
    return figs, tabs, eqs


def emit_index(doc, title, label, items):
    h1(doc, title)
    for no, ttl in items:
        p = para(doc, body=False)
        set_first_line_indent(p, Pt(0))
        p.paragraph_format.left_indent = Pt(12)
        add_rich(p, f'{label}{no}  {ttl}', size=BODY_SIZE)


def emit_references(doc):
    h1(doc, '参考文献')
    md = io.open(os.path.join(ROOT, 'docs', 'REFERENCES_GB7714.md'),
                 encoding='utf-8').read()
    n = 0
    sty = S_REF if S_REF in style_names(doc) else None
    for ln in md.split('\n'):
        s = ln.strip()
        if not re.match(r'^\[\d+\]', s):
            continue
        p = para(doc, sty, body=False)
        clear_char_indent(p)
        pf = p.paragraph_format
        set_first_line_indent(p, -REF_HANG)
        pf.left_indent = REF_HANG
        pf.line_spacing = REF_LINE
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # 条目内可能含行内公式（如 PointCNN 的 $\mathcal{X}$-Transformed），
        # 必须走 OMML，否则 LaTeX 源码直接漏进正文。此处刻意不整体调
        # add_rich：文献条目里的 `*`/`_` 不应被当作 markdown 强调。
        for seg in re.split(r'(\$[^$\n]+\$)', s):
            if not seg:
                continue
            if seg.startswith('$') and seg.endswith('$') and len(seg) > 2:
                append_omml(p, seg[1:-1], display=False)
            else:
                set_run_font(p.add_run(seg), EA_FONT, EN_FONT, REF_SIZE)
        n += 1
    return n


def emit_ack(doc):
    h1(doc, '致谢')
    p0 = os.path.join(CH, 'acknowledgements.md')
    if not os.path.exists(p0):
        return 0
    n = 0
    for ln in io.open(p0, encoding='utf-8').read().split('\n'):
        s = ln.strip()
        if not s or s.startswith('#') or s.startswith('>'):
            continue
        p = para(doc)
        add_rich(p, s, ea='仿宋', en='仿宋', size=BODY_SIZE)
        n += 1
    return n


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = new_document()
    stats = Counter()
    dedup = []

    clear_hf(doc.sections[0])
    set_page_numbering(doc.sections[0], 'decimal', 1)
    cover_zh(doc)
    pbreak(doc)
    cover_en(doc)
    pbreak(doc)
    declaration(doc)

    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    fill_hf(sec2, doc)
    set_page_numbering(sec2, 'lowerRoman', 1)
    emit_abstract(doc, 'abstract_zh.md', False, stats)
    pbreak(doc)
    emit_abstract(doc, 'abstract_en.md', True, stats)
    pbreak(doc)
    h1(doc, '目录')
    p = para(doc)
    set_first_line_indent(p, Pt(0))
    add_field(p, ' TOC \\o "1-3" \\h \\z \\u ')
    pbreak(doc)
    figs, tabs, eqs = collect_assets()
    emit_index(doc, '插图索引', '图', figs)
    pbreak(doc)
    emit_index(doc, '插表索引', '表', tabs)

    sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
    fill_hf(sec3, doc)
    set_page_numbering(sec3, 'decimal', 1)
    for idx, (fn, title) in enumerate(CHAPTERS):
        emit_chapter(doc, os.path.join(CH, fn), title, stats,
                     dedup_log=dedup)
        if idx < len(CHAPTERS) - 1:
            pbreak(doc)
    pbreak(doc)
    nref = emit_references(doc)
    pbreak(doc)
    nack = emit_ack(doc)

    out = os.path.join(OUT_DIR, 'GAAW_thesis_v3.docx')
    doc.save(out)

    # ---- 导出自检（用户硬指标）----
    with zipfile.ZipFile(out) as z:
        xml = z.read('word/document.xml').decode('utf-8')
    nomath = xml.count('<m:oMath>')

    # Word 域代码（`TOC \o "1-3" \h \z \u`）本身是 Word 语法要求的反斜杠，
    # 删掉目录就失效，故不能计入"内容区残留"。此处刻意分开计数并把域代码
    # 白名单化：既守住内容区必须为 0 的硬指标，也不把域代码悄悄抹掉。
    fld = re.findall(r'<w:instrText[^>]*>([^<]*)</w:instrText>', xml)
    nbs_fld = sum(s.count('\\') for s in fld)
    content = re.sub(r'<w:instrText[^>]*>[^<]*</w:instrText>', '', xml)
    nbs = content.count('\\')
    FLD_EXPECT = 4  # 单个 TOC 域：\o \h \z \u

    print('已生成:', os.path.relpath(out, ROOT))
    print(f'  体积 {os.path.getsize(out) / 1024 / 1024:.2f} MB')
    print()
    print('=== 导出自检 ===')
    print(f'  内容区反斜杠计数: {nbs}   (硬指标: 0)')
    print(f'  域代码反斜杠计数: {nbs_fld}   '
          f'(白名单 TOC 开关，期望 {FLD_EXPECT})')
    print(f'  <m:oMath> 计数:          {nomath}   (硬指标: >= 公式数)')
    print(f'  独立公式(eq):            {stats["eq"]}')
    print()
    print('=== 写入统计 ===')
    for k in ('chapter', 'h2', 'h3', 'body', 'list', 'tab', 'tab_bare',
              'fig', 'fig_missing', 'eq', 'note', 'code', 'abstract',
              'fig_dedup'):
        print(f'  {k:14s} {stats[k]}')
    print(f'  references     {nref}')
    print(f'  ack_paras      {nack}')
    print(f'  索引: 图 {len(figs)} 表 {len(tabs)} 式 {len(eqs)}')
    print()
    if dedup:
        print('=== 图题去重位置 ===')
        for chap, no, act, txt in dedup:
            print(f'  {chap}  图{no}  {act}  {txt}')
    else:
        print('图题去重: 0 处（正文说明行均为分析性内容，已保留）')
    ok = (nbs == 0 and nbs_fld == FLD_EXPECT and nomath >= stats['eq']
          and stats['fig_missing'] == 0)
    if not ok:
        print()
        print('!! 自检未通过：内容区反斜杠须为 0，域代码反斜杠须等于白名单值，'
              'OMML 数须不少于独立公式数，图缺失须为 0。')
    return 0 if ok else 1


if __name__ == '__main__':
    sys.exit(main())
