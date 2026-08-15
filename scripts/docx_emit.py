# -*- coding: utf-8 -*-
"""Markdown 解析与章节写入（被 build_docx_tongji.py 导入）。

拆成独立模块的原因：PowerShell 下无法用 heredoc 往已有脚本追加长代码，
且解析逻辑本身独立可测。
"""
import io
import os
import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_docx_tongji import (IMG_BASES, ROOT, add_caption, add_equation,
                               add_image, add_md_table, add_rich,
                               latex_to_text, set_para, set_run_font)

# 图声明表的表头有三种写法（实测确认，首版只认第三种导致 13 张图漏插）：
#   ch1/ch2: | 编号 | 标题 | 类型 | 状态 |            （路径在「状态」列，前缀「已生成」）
#   ch3    : | 编号 | 标题 | 类型 | 位置 | 状态 |      （5 列，路径在末列）
#   ch5/ch6: | 图编号 | 标题 | 数据来源 | 落盘路径 |
# 统一用「表头首列是 编号/图编号，且该表内存在 `图 N.M` 行」来判定。
FIG_DECL_HDR = re.compile(r'^\|\s*(?:图编号|编号)\s*\|')
FIG_ROW = re.compile(r'^图\s*(\d+\.\d+)')
TAB_TITLE = re.compile(r'^\*\*表\s*(\d+\.\d+)\s+(.+?)\*\*\s*$')
H1 = re.compile(r'^#\s+(.+)$')
H2 = re.compile(r'^##\s+(\d+\.\d+)\s+(.+)$')
H3 = re.compile(r'^###\s+(\d+\.\d+\.\d+)\s+(.+)$')
SUBHEAD = re.compile(r'^\*\*(\d+\.\s*[^*]+)\*\*\s*$')


def resolve_img(rel):
    """图路径可能相对仓库根或 docs/，两处都试。"""
    for b in IMG_BASES:
        p = os.path.join(ROOT, b, rel) if b else os.path.join(ROOT, rel)
        if os.path.exists(p):
            return p
    return None


def parse_table_block(lines, i):
    """从第 i 行起收集 markdown 表格，跳过分隔行，返回 (rows, next_i)。"""
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith('|'):
        cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
        if not all(re.fullmatch(r':?-{2,}:?', c) for c in cells if c):
            rows.append(cells)
        i += 1
    return rows, i


def emit_chapter(doc, path, chapter_title, stats):
    """把一个章节 md 写入 doc，按 FORMAT_TONGJI 的字体表逐类映射。"""
    lines = io.open(path, encoding='utf-8').read().split('\n')

    p = doc.add_paragraph()
    set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_pt=0,
             space_before=0, space_after=18)
    set_run_font(p.add_run(chapter_title), '黑体', '黑体', 16, True)
    stats['chapter'] += 1

    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s or H1.match(s) or s.startswith('>') or s.startswith('---'):
            i += 1
            continue

        m = H2.match(s)
        if m:
            p = doc.add_paragraph()
            set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_pt=0,
                     space_before=12, space_after=6)
            set_run_font(p.add_run(f'{m.group(1)} {m.group(2)}'),
                         '黑体', '黑体', 15, False)
            stats['h2'] += 1
            i += 1
            continue

        m = H3.match(s)
        if m:
            p = doc.add_paragraph()
            set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_pt=0,
                     space_before=10, space_after=4)
            set_run_font(p.add_run(f'{m.group(1)} {m.group(2)}'),
                         '黑体', '黑体', 14, False)
            stats['h3'] += 1
            i += 1
            continue

        if s == '$$':
            body, tag = [], None
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                t2 = lines[i].strip()
                mt = re.search(r'\\tag\{([\d.]+)\}', t2)
                if mt:
                    tag = mt.group(1)
                    t2 = re.sub(r'\\tag\{[\d.]+\}', '', t2).strip()
                if t2:
                    body.append(t2)
                i += 1
            i += 1
            add_equation(doc, ' '.join(body), tag)
            stats['eq'] += 1
            continue

        # 图声明表 -> 真正插图 + 图题（图题在图下方，规范 §2.2）
        # 判定需两个条件同时满足：表头首列是「编号/图编号」，且表内确有
        # 「图 N.M」数据行。只看表头会把普通表格（如表 1.1 也有「编号」列）
        # 误判为图声明表。
        if FIG_DECL_HDR.match(s):
            rows, ni = parse_table_block(lines, i)
            fig_rows = [r for r in rows[1:]
                        if r and FIG_ROW.match(r[0].strip())]
            if not fig_rows:
                # 不是图声明表，按普通表格处理
                add_md_table(doc, rows)
                stats['tab_bare'] += 1
                i = ni
                continue
            for row in fig_rows:
                no = re.search(r'(\d+\.\d+)', row[0]).group(1)
                title = row[1] if len(row) > 1 else ''
                png = None
                for c in row:
                    mm = re.search(r'`([^`]+\.png)`', c)
                    if mm:
                        png = mm.group(1)
                        break
                fp = resolve_img(png) if png else None
                if fp:
                    add_image(doc, fp)
                    add_caption(doc, f'图{no} {title}')
                    stats['fig'] += 1
                else:
                    add_caption(doc, f'图{no} {title}（图片缺失）')
                    stats['fig_missing'] += 1
            i = ni
            continue

        # 表题在表上方（规范 §2.2）
        m = TAB_TITLE.match(s)
        if m:
            add_caption(doc, f'表{m.group(1)} {m.group(2)}')
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].lstrip().startswith('|'):
                rows, nj = parse_table_block(lines, j)
                add_md_table(doc, rows)
                stats['tab'] += 1
                i = nj
            else:
                i += 1
            continue

        if s.startswith('|'):
            rows, ni = parse_table_block(lines, i)
            add_md_table(doc, rows)
            stats['tab_bare'] += 1
            i = ni
            continue

        if s.startswith('```'):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1
            for b in buf:
                p = doc.add_paragraph()
                set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_pt=14,
                         first_indent_pt=0)
                set_run_font(p.add_run(b), '宋体', 'Consolas', 9)
            stats['code'] += 1
            continue

        mlist = re.match(r'^(\d+\.|[-*])\s+(.+)$', s)
        if mlist:
            p = doc.add_paragraph()
            set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt=20,
                     first_indent_pt=0)
            p.paragraph_format.left_indent = Pt(24)
            prefix = mlist.group(1)
            add_rich(p, f'{"·" if prefix in ("-", "*") else prefix} '
                        f'{mlist.group(2)}')
            stats['list'] += 1
            i += 1
            continue

        # 段内小标题：宋体 12pt 左 无首行缩进（规范 [183]）
        if SUBHEAD.match(s):
            p = doc.add_paragraph()
            set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_pt=20,
                     first_indent_pt=0)
            add_rich(p, SUBHEAD.match(s).group(1))
            stats['subhead'] += 1
            i += 1
            continue

        if s.startswith(('表注：', '注：')):
            p = doc.add_paragraph()
            set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent_pt=0,
                     space_after=6)
            add_rich(p, s, size=10.5)
            stats['note'] += 1
            i += 1
            continue

        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_pt=20,
                 first_indent_pt=24)
        add_rich(p, s)
        stats['body'] += 1
        i += 1
