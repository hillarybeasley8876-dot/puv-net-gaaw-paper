# -*- coding: utf-8 -*-
"""以参考模板为基底、公式走 OMML 的 docx 生成器（替换纯文本降级方案）。

用户 6 条意见中最严重的一条：公式做纯文本降级后字符本身就是错的
（`pii=₁N`、`p\pi(1)`、`\mathcalP`），格式刷无法修正。本模块改为：
  · 公式一律经 scripts/latex2omml.py 转成 Word 原生 OMML；
  · 文档以 reference/TEMPLATE_tongji_ref.docx 为基底（保留其 styles.xml），
    正文样式直接套用模板的 Normal / Heading 1-3 / Caption / List Paragraph，
    而不是手工重建 —— 这才是「格式刷」的工程化做法。

模板实测规格（scripts/_tmp/probe_template.py）：
  页面 A4 210×297，上下 25.4mm，左右 31.7mm，页眉 20mm，页脚 15mm
  Normal        宋体 10.5pt
  Heading 1     黑体 16pt 粗
  Heading 2     15pt
  Heading 3     黑体 14pt
  Caption       居中（图表题共用）
  List Paragraph 参考文献条目
  公式样式      MTDisplayEquation（Times New Roman 12pt）与「公式」（12pt 居中）
"""
import io
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from latex2omml import latex_to_omml          # noqa: E402

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CH = os.path.join(ROOT, 'docs', 'chapters')
OUT_DIR = os.path.join(ROOT, 'outputs', 'thesis')
TEMPLATE = os.path.join(ROOT, 'reference', 'TEMPLATE_tongji_ref.docx')

TITLE_ZH = '基于梯度自适应对抗约束的点云上采样方法研究'
TITLE_EN = ('Research on Point Cloud Upsampling with '
            'Gradient-Adaptive Adversarial Constraints')

CHAPTERS = [
    ('ch1_introduction.md', '第1章 绪论'),
    ('ch2_related_work.md', '第2章 文献综述'),
    ('ch3_baseline.md', '第3章 基线模型复现与瓶颈诊断'),
    ('ch4_design.md', '第4章 模型假设与研究设计'),
    ('ch5_mechanism.md', '第5章 梯度自适应对抗权重机制'),
    ('ch6_experiments.md', '第6章 实验结果与边界分析'),
    ('ch7_conclusion.md', '第7章 结论与展望'),
]
IMG_BASES = ['', 'docs']

# 模板中可用的样式名（缺失时回退到 None = Normal）
S_H1 = 'Heading 1'
S_H2 = 'Heading 2'
S_H3 = 'Heading 3'
S_CAP = 'Caption'
S_REF = 'List Paragraph'
S_EQ = 'MTDisplayEquation'

# ---------------------------------------------------------------------------
# 任务 B：模板实测规格（scripts/_tmp/probe_tpl_pf.py）
#
# 关键认知：模板的 styles.xml 里 Normal 写的是 1.5 倍行距 + 首行 22pt，
# 但模板**正文实际段落**（690 段中 449 段，占 65%）用的是直接段落格式覆盖：
#   行距 20pt 固定值、首行缩进 24pt、两端对齐、run 字号 12pt（4393 个 run）。
# 论文成稿必须跟随模板的**实际惯例**而非其样式默认值，故此处显式写入实测值，
# 不依赖样式继承。
# ---------------------------------------------------------------------------
BODY_SIZE = 12.0          # 小四，模板正文 run 实测众数（4393 次）
BODY_LINE = Pt(20)        # 行距固定值 20pt，模板正文实测
BODY_INDENT = Pt(24)      # 首行缩进 24pt = 2 字符 @12pt
CAP_SIZE = 10.5           # 图表题五号
CAP_INDENT = Pt(21)       # Caption 实测首行 21pt（19/21 段）
REF_SIZE = 10.5           # 参考文献五号
REF_HANG = Pt(19.8)       # List Paragraph 实测悬挂 -19.8 / 左 19.8（72/76 段）
REF_LINE = Pt(16)         # 参考文献行距 16pt，模板实测
TAB_SIZE = 9.0            # 表格正文
NOTE_SIZE = 9.0           # 表注
MONO_SIZE = 10.5          # 行内等宽（`...`）：固定五号，不做 size-0.5 机械缩放
EA_FONT = '宋体'           # 中文
EN_FONT = 'Times New Roman'
# 模板实测：Normal 段落里含拉丁字母的 run 有 715 个为 Times New Roman、
# List Paragraph 220 个亦为 Times New Roman（Arial 仅 1 个）。
# 教训：先前把西文也设成宋体，PDF 渲染出等宽风格的英文（见 p140 参考文献），
# 与模板惯例不符。西文默认必须是 Times New Roman。


def set_first_line_indent(p, pt):
    """设置首行缩进，并同时清掉字符单位缩进。

    必须用本函数替代裸 `paragraph_format.first_line_indent = Pt(x)`：
    OOXML 的 `firstLineChars` 优先于 `firstLine`，而 python-docx 只写后者。
    实测教训（图索引/表索引 38 段）：写了 `first_line_indent = Pt(0)` 却
    仍继承 Normal 的 `firstLineChars="200"`，Word 实际缩进 2 字符 = 24pt，
    与正文段无法区分。这类漏洞在本项目出现过三次（公式段、索引段、Caption），
    故收敛为单一入口。
    """
    clear_char_indent(p)
    p.paragraph_format.first_line_indent = pt
    return p


def apply_body_format(p):
    """正文段落格式：模板实测惯例（行距 20pt、首行 24pt、两端对齐）。"""
    pf = p.paragraph_format
    pf.line_spacing = BODY_LINE
    set_first_line_indent(p, BODY_INDENT)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def new_document():
    """以模板为基底建文档，清空正文但保留 styles / numbering / 页面设置。"""
    doc = Document(TEMPLATE)
    body = doc.element.body
    # 删除除最后一个 sectPr 之外的全部内容
    sect = body.find(qn('w:sectPr'))
    for child in list(body):
        if child is not sect:
            body.remove(child)
    return doc


def style_names(doc):
    return {s.name for s in doc.styles}


def suppress_auto_numbering(p):
    """在段落级关闭 Heading 样式自带的自动编号。

    模板的 Heading 1/2/3 都挂了 `w:numPr`，Word 会自动生成「第10章」
    「10.1」「10.1.1」这类编号，与本文手写的章节号叠加成
    「第10章 第 5 章 …」「10.1 5.1 …」（实测 p088 渲染确认）。

    本文刻意保留**手写章节号**而不改用 Word 自动编号：全文的交叉引用
    （节 169 / 图 16 / 表 22 个锚点）与图表编号体系已与手写章号严格绑定，
    交给 Word 重新生成会与正文引用失配。

    做法是在段落 pPr 内写入 `<w:numPr><w:ilvl val=0/><w:numId val=0/>`，
    numId=0 是「无列表」，按 OOXML 规范可覆盖样式继承的编号。
    """
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn('w:numPr'))
    if old is not None:
        pPr.remove(old)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '0')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    return p


HEADING_STYLES = {S_H1, S_H2, S_H3}


def clear_char_indent(p):
    """清除段落继承的字符单位缩进（`w:firstLineChars` / `w:leftChars`）。

    坑：OOXML 的 `firstLineChars`（字符单位）**优先于** `firstLine`（缇），
    而 python-docx 的 `paragraph_format.first_line_indent` 只写 `firstLine`。
    模板 MTDisplayEquation 样式带 `firstLineChars="1900"`（19 字符），
    仅设 first_line_indent=Pt(0) 完全压不住：实测公式被推到版心右侧
    x0=386pt 起笔然后折行，改了三轮参数毫无变化（见 p088 渲染）。
    必须显式删掉这两个属性，或把它们置 0。
    """
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    for attr in ('firstLineChars', 'leftChars', 'rightChars',
                 'hangingChars', 'startChars', 'endChars'):
        ind.set(qn('w:' + attr), '0')
    return p


def para(doc, style=None, body=None):
    """新建段落。

    body=None 时按「无样式即正文」推断：未指定样式的段落一律套模板实测的
    正文格式（行距 20pt / 首行 24pt / 两端对齐）。显式传 body=False 可跳过，
    用于表格单元格、表注等自定格式段。

    标题样式段落一律抑制自动编号（见 suppress_auto_numbering）。
    """
    p = doc.add_paragraph()
    if style and style in style_names(doc):
        p.style = doc.styles[style]
        if style in HEADING_STYLES:
            suppress_auto_numbering(p)
    if body is None:
        body = style is None
    if body:
        apply_body_format(p)
    return p


def set_run_font(run, ea, en, size=None, bold=None):
    """设置中西文字体。中文必须走 w:eastAsia。"""
    if en:
        run.font.name = en
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    if ea:
        rFonts.set(qn('w:eastAsia'), ea)
    if en:
        rFonts.set(qn('w:ascii'), en)
        rFonts.set(qn('w:hAnsi'), en)


def append_omml(p, latex, display=False):
    """把 LaTeX 作为 OMML 追加到段落 p。"""
    xml = latex_to_omml(latex, display=display)
    el = parse_xml(f'<w:p {nsdecls("w", "m")}>{xml}</w:p>')
    for child in list(el):
        p._p.append(child)


INLINE = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\$[^$\n]+\$)')
CITE_RE = re.compile(r'(\[\d+(?:,\d+)*\])')


def unescape_text(s):
    for a, b in (('\\%', '%'), ('\\&', '&'), ('\\#', '#'),
                 ('\\_', '_'), ('\\$', '$'), ('\\{', '{'), ('\\}', '}')):
        s = s.replace(a, b)
    return s


def add_rich(p, text, ea=EA_FONT, en=EN_FONT, size=None,
             base_bold=False):
    """行内富文本：粗体、等宽、**公式走 OMML**、引用编号上标。

    注意：必须递归处理，否则「粗体内嵌行内公式」（`**...$x$...**`）会被
    粗体分支整块吞掉，其中的 $...$ 退化为纯文本（实测源稿 70 处）。
    """
    def emit_plain(txt, bold):
        for piece in CITE_RE.split(txt):
            if not piece:
                continue
            r = p.add_run(unescape_text(piece))
            set_run_font(r, ea, en, size, bold)
            if CITE_RE.fullmatch(piece):
                r.font.superscript = True

    def walk(txt, bold):
        for seg in INLINE.split(txt):
            if not seg:
                continue
            if seg.startswith('**') and seg.endswith('**') and len(seg) > 4:
                walk(seg[2:-2], True)
            elif seg.startswith('`') and seg.endswith('`') and len(seg) > 2:
                r = p.add_run(seg[1:-1])
                set_run_font(r, ea, 'Consolas',
                             min(MONO_SIZE, size) if size else MONO_SIZE,
                             bold)
            elif seg.startswith('$') and seg.endswith('$') and len(seg) > 2:
                append_omml(p, seg[1:-1], display=False)
            else:
                emit_plain(seg, bold)

    walk(text, base_bold)
