"""抽同济示例里表格的真实结构：边框类型、对齐、字体，判断是否三线表。"""
from docx import Document
from docx.oxml.ns import qn

DOCX = r"E:\AE-CC托管\puv-net\refs\tongji_template\tongji_example.docx"
doc = Document(DOCX)

print("表格数:", len(doc.tables))
for ti, t in enumerate(doc.tables):
    print("=" * 70)
    print("表 #{}  style={}  rows={} cols={}".format(
        ti, t.style.name if t.style else None, len(t.rows), len(t.columns)))
    # 表级边框设置
    tblPr = t._tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        print("  tblBorders: 无（继承样式）")
    else:
        for b in borders:
            tag = b.tag.split('}')[1]
            print("    border {:8s} val={} sz={} color={}".format(
                tag, b.get(qn('w:val')), b.get(qn('w:sz')), b.get(qn('w:color'))))
    print("  --- 单元格内容 ---")
    for ri, row in enumerate(t.rows):
        cells = []
        for c in row.cells:
            txt = c.text.strip().replace("\n", "/")
            cells.append(txt if txt else "·")
        print("    r{}: {}".format(ri, " | ".join(cells)))
    # 首行单元格字体
    print("  --- 首行字体 ---")
    for c in t.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                if r.text.strip():
                    rPr = r._element.rPr
                    ea = rPr.rFonts.get(qn('w:eastAsia')) if (rPr is not None and rPr.rFonts is not None) else None
                    print("      '{}' ea={} sz={} bold={} align={}".format(
                        r.text.strip()[:20], ea,
                        round(r.font.size.pt, 1) if r.font.size else None,
                        r.font.bold, p.paragraph_format.alignment))
                    break
            break
