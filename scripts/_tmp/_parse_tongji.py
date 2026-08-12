"""用 Word COM 把 .doc 转 docx，再用 python-docx 抽格式规范。只读+另存，不改原文件。"""
import os, sys, json

SRC = r"C:\Users\nieyi\.minimax\v2\assets\2026\08\12\22-46-49-830-asset_20260812-224649-830_214ed9eec700_185f11c0-同济大学研究生学位论文写作示例（参考）.doc"
OUTDIR = r"E:\AE-CC托管\puv-net\refs\tongji_template"
os.makedirs(OUTDIR, exist_ok=True)
DOCX = os.path.join(OUTDIR, "tongji_example.docx")

if not os.path.exists(DOCX):
    import win32com.client as win32
    w = win32.Dispatch("Word.Application")
    w.Visible = False
    try:
        d = w.Documents.Open(SRC, ReadOnly=True)
        d.SaveAs2(DOCX, FileFormat=16)  # 16 = wdFormatDocumentDefault (.docx)
        d.Close(False)
    finally:
        w.Quit()
    print("converted ->", DOCX)
else:
    print("already exists ->", DOCX)

from docx import Document
from docx.shared import Pt

doc = Document(DOCX)
print("=" * 78)
print("段落总数:", len(doc.paragraphs), " 表格数:", len(doc.tables))
print("=" * 78)

# 1) 样式清单
print("\n--- 文档中实际用到的段落样式 & 计数 ---")
from collections import Counter, OrderedDict
c = Counter(p.style.name for p in doc.paragraphs if p.text.strip())
for k, v in c.most_common():
    print("  {:34s} {}".format(k, v))

# 2) 每个用到的样式的字体/段落格式
print("\n--- 各样式格式细节 ---")
for name in c:
    try:
        st = doc.styles[name]
    except KeyError:
        continue
    f = st.font
    pf = st.paragraph_format
    def sz(x):
        return round(x.pt, 1) if x is not None else None
    def ln(x):
        if x is None:
            return None
        return round(x.pt, 1) if hasattr(x, "pt") else x
    print("  [{}]".format(name))
    print("     font={} size={} bold={} eastasia={}".format(
        f.name, sz(f.size), f.bold,
        f.element.rPr.rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
        if f.element.rPr is not None and f.element.rPr.rFonts is not None else None))
    print("     align={} line_spacing={} space_before={} space_after={} first_indent={}".format(
        pf.alignment, ln(pf.line_spacing), sz(pf.space_before), sz(pf.space_after),
        sz(pf.first_line_indent)))
